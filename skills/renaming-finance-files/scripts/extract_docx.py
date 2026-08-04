#!/usr/bin/env python3
"""Extract text content from DOCX files for finance file processing."""
import sys
from docx import Document

def extract_docx(filepath):
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    print("\n".join(lines))

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print(f"Usage: {sys.argv[0]} <filepath.docx>", file=sys.stderr)
        sys.exit(1)
    extract_docx(sys.argv[1])
