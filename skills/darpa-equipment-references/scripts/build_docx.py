#!/usr/bin/env python3
"""Build the DARPA Equipment List .docx from parsed input items.

Workflow:
    1. Load parsed items (parse_input.py output).
    2. Optionally load existing-entries map (extract_existing.py output) and
       merge: for each input item, if an existing entry matches by link or
       (product/title), reuse its `product` (full product name) and
       `explanation`. Otherwise leave them blank — caller is expected to fill
       these in before final build.
    3. Open template docx (the existing Equipment List.docx by default) and
       rewrite the body paragraphs from "EQ001 ..." onward. The preamble
       ("Additional Budget Information", "Equipment List") and all
       header/footer/image content are preserved.

Usage:
    build_docx.py --items items.json --template <template.docx> --out <out.docx>
                  [--existing existing.json] [--strict]

`--strict` aborts if any item is missing a `product` or `explanation` instead
of writing the literal placeholder. Use it for the final build.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from copy import deepcopy
from pathlib import Path


PREAMBLE_MARKERS = ("Additional Budget Information", "Equipment List")


def _fmt_money(v: float) -> str:
    # No fractional cents when value is whole, else 2 dp
    if abs(v - round(v)) < 0.005:
        return f"${int(round(v)):,}"
    return f"${v:,.2f}"


def _cost_qty_line(item: dict) -> str:
    cost = float(item.get("cost_per_item") or 0)
    qty = int(item.get("qty") or 1)
    total = float(item.get("total_cost") or cost * qty)
    return f"{_fmt_money(cost)} × {qty} = {_fmt_money(total)}"


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def merge_existing(items: list[dict], existing: list[dict]) -> list[dict]:
    """Best-effort match input items to existing entries by link or title."""
    by_link = {_norm(e.get("link", "").split("?", 1)[0]): e for e in existing if e.get("link")}
    by_title = {_norm(e.get("title", "")): e for e in existing if e.get("title")}

    merged = []
    for it in items:
        link_key = _norm(it.get("link", "").split("?", 1)[0])
        match = by_link.get(link_key)
        if not match:
            match = by_title.get(_norm(it.get("description", "")))
        # Prefer the existing doc's product/explanation only when it has one;
        # otherwise fall back to what items.json already has. Without this,
        # items whose URL doesn't appear in the existing doc would be wiped.
        product = (match.get("product") if match else "") or it.get("product") or it.get("description", "")
        explanation = (match.get("explanation") if match else "") or it.get("explanation", "")
        merged.append({**it, "product": product, "explanation": explanation})
    return merged


def _clear_body_after_preamble(doc) -> int:
    """Delete every paragraph after the 'Equipment List' preamble. Returns the
    index of the last preamble paragraph (next inserts go after it)."""
    body = doc.element.body
    paragraphs = doc.paragraphs

    preamble_end = None
    for i, p in enumerate(paragraphs):
        if p.text.strip().lower().startswith("equipment list"):
            preamble_end = i
            break
    if preamble_end is None:
        # No preamble found — keep just the very first paragraph and clear the rest
        preamble_end = 0

    # Remove every paragraph after preamble_end. Operate on raw XML so we don't
    # disturb any inline images / shapes that live alongside text.
    keep = paragraphs[: preamble_end + 1]
    keep_elems = {id(p._element) for p in keep}

    for child in list(body):
        if child.tag.endswith("}p") and id(child) not in keep_elems:
            body.remove(child)
        elif child.tag.endswith("}tbl"):
            # Drop tables that appear after the preamble (they're equipment-block leftovers).
            # Tables that appear *before* preamble (rare) are kept.
            body.remove(child)
    return preamble_end


def _add_run(p, text: str, *, bold: bool = False, size_pt: int | None = None):
    from docx.shared import Pt
    r = p.add_run(text)
    r.bold = bold
    if size_pt:
        r.font.size = Pt(size_pt)
    return r


_HYPERLINK_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


def _add_hyperlink(paragraph, url: str, text: str, color: str = "0563C1") -> None:
    """Append a real (clickable) Word hyperlink to the given paragraph."""
    from docx.oxml.shared import OxmlElement, qn

    part = paragraph.part
    r_id = part.relate_to(url, _HYPERLINK_REL, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _spacer(doc):
    doc.add_paragraph("")


def _add_page_break_before(paragraph) -> None:
    """Insert a hard page break as the first run of the paragraph so the
    paragraph (and everything after it) starts at the top of a new page."""
    from docx.oxml.shared import OxmlElement, qn
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    paragraph._p.insert(0, r)


def write_entry(doc, item: dict, idx: int, strict: bool) -> None:
    """Append one EQ block to the document, mirroring the original layout."""
    from docx.shared import Pt

    eq_no = f"EQ{idx:03d}"
    title = item.get("description") or item.get("product") or "Untitled"
    product = item.get("product") or item.get("description") or ""
    vendor = item.get("vendor") or ""
    cost_qty = _cost_qty_line(item)
    explanation = item.get("explanation") or ""
    link = item.get("link") or ""

    if strict:
        missing = []
        if not product:
            missing.append("product")
        if not explanation:
            missing.append("explanation")
        if missing:
            sys.exit(f"ERROR: {eq_no} ({title!r}) missing: {', '.join(missing)}. "
                     f"Re-run without --strict, fill the items JSON, then re-build with --strict.")

    if not explanation:
        explanation = "[EXPLANATION NEEDED — fill in items JSON]"

    # Heading (bold). Every EQ entry starts on its own page (except EQ001 —
    # the preamble already sits at the top, so we don't want a leading blank
    # page).
    p = doc.add_paragraph()
    _add_run(p, f"{eq_no} - {title}", bold=True)
    if idx > 1:
        _add_page_break_before(p)

    _spacer(doc)

    # Body block — Product / Vendor / Cost & Quantity
    p = doc.add_paragraph()
    _add_run(p, "Product: ", bold=True)
    _add_run(p, product)
    p.add_run("\n")
    _add_run(p, "Vendor: ", bold=True)
    _add_run(p, vendor)
    p.add_run("\n")
    _add_run(p, "Cost & Quantity: ", bold=True)
    _add_run(p, cost_qty)

    # Explanation
    p = doc.add_paragraph()
    _add_run(p, "Explanation: ", bold=True)
    _add_run(p, explanation)

    # Link — real clickable hyperlink
    p = doc.add_paragraph()
    _add_run(p, "Link: ", bold=True)
    if link:
        _add_hyperlink(p, link, link)
    else:
        _add_run(p, "")

    # Screenshot of the product page (if cached)
    screenshot = item.get("screenshot")
    if screenshot and Path(screenshot).exists():
        from docx.shared import Inches
        p = doc.add_paragraph()
        run = p.add_run()
        try:
            run.add_picture(screenshot, width=Inches(5.5))
        except Exception as e:
            print(f"WARN: failed to embed screenshot for {eq_no}: {e}", file=sys.stderr)

    # Hard page break before the next EQ heading separates items — no
    # trailing spacers needed.


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--items", type=Path, required=True)
    ap.add_argument("--template", type=Path, required=True,
                    help="Existing docx used as style/preamble template (will not be modified).")
    ap.add_argument("--existing", type=Path, default=None,
                    help="Optional extract_existing.py output to preserve explanations.")
    ap.add_argument("--out", type=Path, required=True)
    ap.add_argument("--strict", action="store_true",
                    help="Abort if any item is missing product or explanation.")
    args = ap.parse_args()

    try:
        from docx import Document
    except ImportError:
        sys.exit("ERROR: python-docx not installed. pip install python-docx")

    if not args.template.exists():
        sys.exit(f"ERROR: template not found: {args.template}")
    items = json.loads(args.items.read_text(encoding="utf-8"))
    existing = []
    if args.existing and args.existing.exists():
        existing = json.loads(args.existing.read_text(encoding="utf-8"))

    if existing:
        items = merge_existing(items, existing)
    else:
        # Ensure expected keys exist
        for it in items:
            it.setdefault("product", it.get("description", ""))
            it.setdefault("explanation", "")

    doc = Document(str(args.template))
    _clear_body_after_preamble(doc)

    # Add a couple of spacer paragraphs after preamble to match original spacing
    _spacer(doc)
    _spacer(doc)

    for i, it in enumerate(items, 1):
        write_entry(doc, it, i, args.strict)

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out))
    print(f"Wrote {len(items)} equipment entries → {args.out}")


if __name__ == "__main__":
    main()
